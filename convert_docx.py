import docx
import sys

def docx_to_text(filename):
    try:
        doc = docx.Document(filename)
        text = []
        for p in doc.paragraphs:
            text.append(p.text)
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    row_text.append(cell.text.replace('\n', ' '))
                text.append(' | '.join(row_text))
        return '\n'.join(text)
    except Exception as e:
        return str(e)

with open('theme.md', 'w', encoding='utf-8') as f:
    f.write(docx_to_text('Introducción a la Ciencia de datos.docx'))

with open('syllabus.md', 'w', encoding='utf-8') as f:
    f.write(docx_to_text('Sylabus INTRODUCCIÓN A LA CIENCIA DE DATOS.docx'))

with open('template.md', 'w', encoding='utf-8') as f:
    f.write(docx_to_text('Guia-de-Aprendizaje_UNAB1.docx'))
