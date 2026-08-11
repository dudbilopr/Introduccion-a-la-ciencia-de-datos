import docx
from docx.shared import Pt
import re

def main():
    doc_path = 'Guia-de-Aprendizaje_UNAB1.docx'
    out_path = 'Guia_de_Aprendizaje_Final.docx'
    
    try:
        doc = docx.Document(doc_path)
    except Exception as e:
        print("Error loading document:", e)
        return

    # Replace simple text in paragraphs
    for p in doc.paragraphs:
        if 'Profesor (a):' in p.text:
            p.text = p.text.replace('Profesor (a):', 'Profesor (a): Dudbil Olvasada Pabon Riaño')
        if 'Facultad:' in p.text:
            p.text = p.text.replace('Facultad:', 'Facultad: Ciencias Básicas')
        if 'Programa:' in p.text:
            p.text = p.text.replace('Programa:', 'Programa: Ciencia de Datos')

        # Justificación / Resumen
        if 'Presente en uno o dos párrafos información que le permita' in p.text:
            p.text = "En este curso se desarrollan los conceptos básicos del análisis de conglomerados y luego el estudio de un portafolio de metodologías, algoritmos y aplicaciones de agrupamiento típicos. Esto incluye métodos de partición como k-means, métodos jerárquicos como BIRCH y métodos basados ​​en densidad como DBSCAN/OPTICS. Además, se tratan métodos para la validación de agrupaciones y la evaluación de la calidad de la agrupación. Finalmente, se aplicarán ejemplos de análisis de conglomerados en diversas situaciones."
        
        # Competencias
        if 'Competencias Disciplinares: copie y pegue del Syllabus' in p.text:
            p.text = (
                "Competencias Disciplinares:\n"
                "1. Interpreta los modelos de agrupamiento como herramientas de clasificación y reconocimiento de patrones.\n"
                "2. Diseña diferentes tipos de agrupamiento según los diversos enfoques.\n"
                "3. Valida los modelos de conglomerados aplicados a situaciones prácticas de acuerdo con criterios establecidos."
            )
            
        # EHLE
        if '3. ¿Es un curso tipo EHLE?' in p.text:
            p.text = '3. ¿Es un curso tipo EHLE? Sí___ No_X_'
        if 'En caso afirmativo, explíqueles' in p.text:
            p.text = ''
            
        # Inglés
        if '4. ¿Este curso integra recursos o actividades en inglés?' in p.text:
            p.text = '4. ¿Este curso integra recursos o actividades en inglés?  Sí_X_ No___'
        if 'En caso de integrar el inglés explíqueles' in p.text:
            p.text = 'Se implementará la estrategia EMI (English Medium Instruction) en inglés durante 10 minutos en cada clase para fomentar el bilingüismo y la familiarización con la terminología técnica.'
            
        # Acuerdos
        if 'A continuación, se presentan algunos ítems sobre acuerdos particulares' in p.text:
            p.text = (
                "Acuerdos Particulares:\n"
                "- Horario: Las clases inician de manera puntual todos los martes de 7:00 a 10:00 de la mañana.\n"
                "- Llegar temprano: Se exige puntualidad en el ingreso a la sesión.\n"
                "- Descansos: Siempre se dejará un tiempo de descanso durante la sesión.\n"
                "- Estrategia EMI: Se realizará una actividad en inglés de 10 minutos.\n"
                "- Respeto: El respeto mutuo es obligatorio en todo momento del desarrollo de la clase.\n"
                "- Comunicación: Se manejará un grupo en Gmail (un espacio) para mantener la comunicación constante con los estudiantes."
            )
        # Remove placeholder bullet points for acuerdos
        if any(x in p.text for x in ["Hora de inicio y finalización", "Uso del celular en clase", "Maneras de intervención", "Descansos.", "Entrega de actividades", "Orientaciones para trabajo en grupo"]):
            p.text = ""

    # Tables
    # Identify tables by their headers
    for table in doc.tables:
        if len(table.rows) == 0:
            continue
        first_row_text = ' '.join(cell.text for cell in table.rows[0].cells)
        
        # Course Info Table (Nombre del curso, Periodo, etc.)
        if 'Nombre del curso / módulo' in first_row_text or 'Nombre del curso' in first_row_text:
            # Usually cell 0 is label, cell 1 is value, etc. But let's just replace in all cells
            for row in table.rows:
                for cell in row.cells:
                    if 'Nombre del curso / módulo:' in cell.text:
                        cell.text = 'Nombre del curso / módulo: INTROD A LA CIENCIA DE DATOS (Grupo: B13-AINF, Código: 63448)'
                    if 'Horas  Teóricas' in cell.text:
                        cell.text = cell.text.replace('Horas  Teóricas', 'Horas Teóricas: 3')
                    if 'Horas  Prácticas' in cell.text:
                        cell.text = cell.text.replace('Horas  Prácticas', 'Horas Prácticas: 0')
                    if 'Horas Teórico  Prácticas' in cell.text:
                        cell.text = cell.text.replace('Horas Teórico  Prácticas', 'Horas Teórico Prácticas: 3')
                    if 'Horas de Trabajo Independiente' in cell.text:
                        cell.text = cell.text.replace('Horas de Trabajo Independiente', 'Horas de Trabajo Independiente: 6')
                    if 'Créditos' in cell.text and ':' not in cell.text:
                        cell.text = 'Créditos: 3'
        
        # Schedule Table (Semana o sesión, Contenidos, Actividades)
        if 'Semana o sesión' in first_row_text:
            # We will populate the table with the 16 weeks
            themes = [
                ("Semana 1", "Introducción a la Ciencia de Datos. Definición, el ciclo de vida de los datos y roles en el equipo.", "Clase magistral, discusión. Actividad de familiarización."),
                ("Semana 2", "Herramientas del Ecosistema. Instalación y uso de Python (Pandas, NumPy) o R, y entornos como Jupyter Notebooks.", "Laboratorio de instalación y primeros pasos en código."),
                ("Semana 3", "Adquisición de Datos. Tipos de datos (estructurados vs. no estructurados) y formatos comunes como CSV, JSON y XML.", "Laboratorio de carga de datos."),
                ("Semana 4", "Limpieza y Preprocesamiento. Manejo de valores faltantes, detección de valores atípicos (outliers) y normalización de datos.", "Laboratorio de preprocesamiento de un dataset real."),
                ("Semana 5", "Estadística Descriptiva. Medidas de tendencia central, variabilidad y posición.", "Ejercicios prácticos de estadística en Python/R."),
                ("Semana 6", "Visualización de Datos. Creación de histogramas, diagramas de caja (boxplots) y gráficos de dispersión para identificar patrones.", "Taller de visualización interactiva."),
                ("Semana 7", "Análisis Exploratorio de Datos (EDA). Técnicas para formular hipótesis y encontrar relaciones entre variables.", "Taller de EDA."),
                ("Semana 8", "Examen Parcial I. Evaluación de conceptos de preprocesamiento y análisis descriptivo.", "Evaluación escrita/práctica."),
                ("Semana 9", "Inferencia Estadística. Población y muestra, intervalos de confianza y pruebas de hipótesis.", "Resolución de problemas de inferencia."),
                ("Semana 10", "Introducción al Aprendizaje Supervisado (Regresión). Modelado de relaciones continuas mediante regresión lineal simple y múltiple.", "Laboratorio de modelado predictivo."),
                ("Semana 11", "Aprendizaje Supervisado (Clasificación). Problemas de respuesta binaria usando regresión logística y árboles de decisión.", "Laboratorio de clasificación."),
                ("Semana 12", "Evaluación de Modelos. Uso de la matriz de confusión, métricas de precisión, sensibilidad (recall) y validación cruzada.", "Taller de evaluación de modelos."),
                ("Semana 13", "Aprendizaje No Supervisado. Descubrimiento de estructuras ocultas mediante agrupamiento (clustering) como el algoritmo K-means.", "Laboratorio de segmentación y agrupamiento."),
                ("Semana 14", "Aplicaciones Especializadas. Introducción breve a series de tiempo, análisis de redes o procesamiento de lenguaje natural (NLP).", "Estudio de caso especializado."),
                ("Semana 15", "Ética y Ciencia de Datos Responsable. Privacidad, sesgo en algoritmos, transparencia y responsabilidad social.", "Debate y análisis crítico de casos de estudio."),
                ("Semana 16", "Examen Parcial II / Proyecto Final. Evaluación del modelado y presentación de resultados finales.", "Sustentación de proyecto final.")
            ]
            
            # Clear existing data rows (keep header)
            while len(table.rows) > 1:
                table._element.remove(table.rows[1]._element)
                
            for week, content, activity in themes:
                row_cells = table.add_row().cells
                row_cells[0].text = week
                row_cells[1].text = content
                row_cells[2].text = activity

        # Evaluation Table
        if 'Actividad de evaluación' in first_row_text:
            evaluations = [
                ("Semana 8", "Examen Parcial I (Conceptos de preprocesamiento, EDA, y distancia)", "RAE 1.1, 1.2, 1.3", "Individual", "20%"),
                ("Semanas 2-7, 9-13", "Laboratorios y Talleres Prácticos", "RAE 2.1, 2.2, 2.3", "Grupal/Individual", "40%"),
                ("Semana 16", "Proyecto Final y Examen Parcial II", "RAE 3.1, 3.2, 3.3", "Grupal", "40%")
            ]
            while len(table.rows) > 1:
                table._element.remove(table.rows[1]._element)
            
            for date, activity, rae, act_type, perc in evaluations:
                row_cells = table.add_row().cells
                row_cells[0].text = date
                row_cells[1].text = activity
                row_cells[2].text = rae
                row_cells[3].text = act_type
                row_cells[4].text = perc

    doc.save(out_path)
    print(f"Document successfully created at {out_path}")

if __name__ == '__main__':
    main()
