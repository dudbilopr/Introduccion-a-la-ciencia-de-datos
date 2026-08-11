import docx

def main():
    doc_path = 'Guia-de-Aprendizaje_UNAB1.docx'
    out_path = 'Guia_de_Aprendizaje_Final.docx'
    
    try:
        doc = docx.Document(doc_path)
    except Exception as e:
        print("Error loading document:", e)
        return

    # Replace simple text in paragraphs
    for p in doc.paragraphs:
        if 'Profesor (a):' in p.text:
            p.text = p.text.replace('Profesor (a):', 'Profesor (a): Dudbil Olvasada Pabon Riaño')
            
        if 'Presentación: ¿qué quiere compartirle' in p.text:
            p.text = "¡Hola! Soy Dudbil Olvasada Pabon Riaño, Magíster en Física y Coordinador del programa de Ciencia de Datos en la UNAB. Mi pasión se encuentra en la intersección entre las ciencias exactas (como la Óptica Cuántica y el Electromagnetismo) y la tecnología de vanguardia. Tengo amplia experiencia en Machine Learning, Big Data, análisis de sentimiento y desarrollo con Python (Django/Flask). Me entusiasma enseñarles no solo a programar, sino a pensar analíticamente para resolver retos complejos usando datos e Inteligencia Artificial."
            
        if 'Enlace a perfil(es):' in p.text:
            p.text = "Google Scholar: https://scholar.google.com/citations?user=HMVbVhcAAAAJ | ORCID: 0000-0001-7239-3497"
            
        if 'Facultad:' in p.text:
            p.text = p.text.replace('Facultad:', 'Facultad: Ciencias Básicas')
        if 'Programa:' in p.text:
            p.text = p.text.replace('Programa:', 'Programa: Ciencia de Datos')

        # Justificación / Resumen
        if 'Presente en uno o dos párrafos información que le permita' in p.text:
            p.text = "En este curso se desarrollan los conceptos básicos del análisis de datos desde una perspectiva de aprendizaje activo y basado en retos. Iniciamos con la preparación y limpieza de datos, avanzamos a la estadística descriptiva e inferencial, y culminamos con modelado predictivo, incluyendo análisis de conglomerados (k-means, BIRCH, DBSCAN). Todo esto se enmarca en la toma de decisiones, la evaluación rigurosa de modelos y la ética profesional. Nuestro objetivo es que el estudiante no solo programe, sino que desarrolle un pensamiento crítico para resolver problemas reales mediante datos."
        
        # Competencias
        if 'Competencias Disciplinares: copie y pegue del Syllabus' in p.text:
            p.text = (
                "Competencias Disciplinares:\n"
                "1. Interpreta los modelos de agrupamiento como herramientas de clasificación y reconocimiento de patrones.\n"
                "2. Diseña diferentes tipos de agrupamiento según los diversos enfoques.\n"
                "3. Valida los modelos de conglomerados aplicados a situaciones prácticas de acuerdo con criterios establecidos."
            )
            
        # EHLE
        if '3. ¿Es un curso tipo EHLE?' in p.text:
            p.text = '3. ¿Es un curso tipo EHLE? Sí___ No_X_'
        if 'En caso afirmativo, explíqueles' in p.text:
            p.text = ''
            
        # Inglés
        if '4. ¿Este curso integra recursos o actividades en inglés?' in p.text:
            p.text = '4. ¿Este curso integra recursos o actividades en inglés?  Sí_X_ No___'
        if 'En caso de integrar el inglés explíqueles' in p.text:
            p.text = 'Se implementará la estrategia EMI (English Medium Instruction) durante 10 minutos en cada clase para fomentar el bilingüismo, debatiendo sobre literatura científica y terminología técnica en el ecosistema global de ciencia de datos.'
            
        # Acuerdos
        if 'A continuación, se presentan algunos ítems sobre acuerdos particulares' in p.text:
            p.text = (
                "Acuerdos Particulares (Consensuados):\n"
                "- Horario y Puntualidad: Las clases inician de manera puntual todos los martes a las 7:00 a.m. y terminan a las 10:00 a.m. Se exige respeto por el tiempo de los demás.\n"
                "- Descansos Cognitivos: Siempre se dejará un tiempo de descanso activo durante la sesión para evitar la fatiga y mejorar la atención sostenida.\n"
                "- Estrategia EMI: Espacio inmersivo en inglés de 10 minutos semanales.\n"
                "- Respeto y Convivencia: Cero tolerancia a la falta de respeto. Fomentamos un ambiente psicológicamente seguro para debatir ideas.\n"
                "- Comunicación y Colaboración: Se manejará un grupo en Gmail (un espacio de chat/comunidad) para consultas asíncronas y apoyo mutuo entre pares."
            )
        # Remove placeholder bullet points for acuerdos
        if any(x in p.text for x in ["Hora de inicio y finalización", "Uso del celular en clase", "Maneras de intervención", "Descansos.", "Entrega de actividades", "Orientaciones para trabajo en grupo"]):
            p.text = ""

    # Tables
    for table in doc.tables:
        if len(table.rows) == 0:
            continue
        first_row_text = ' '.join(cell.text for cell in table.rows[0].cells)
        
        # Course Info Table 
        if 'Nombre del curso / módulo' in first_row_text or 'Nombre del curso' in first_row_text:
            for row in table.rows:
                for cell in row.cells:
                    if 'Nombre del curso / módulo:' in cell.text:
                        cell.text = 'Nombre del curso / módulo: INTROD A LA CIENCIA DE DATOS (Grupo: B13-AINF, Código: 63448)'
                    if 'Horas  Teóricas' in cell.text:
                        cell.text = cell.text.replace('Horas  Teóricas', 'Horas Teóricas: 3')
                    if 'Horas  Prácticas' in cell.text:
                        cell.text = cell.text.replace('Horas  Prácticas', 'Horas Prácticas: 0')
                    if 'Horas Teórico  Prácticas' in cell.text:
                        cell.text = cell.text.replace('Horas Teórico  Prácticas', 'Horas Teórico Prácticas: 3')
                    if 'Horas de Trabajo Independiente' in cell.text:
                        cell.text = cell.text.replace('Horas de Trabajo Independiente', 'Horas de Trabajo Independiente: 6')
                    if 'Créditos' in cell.text and ':' not in cell.text:
                        cell.text = 'Créditos: 3'
        
        # Schedule Table 
        if 'Semana o sesión' in first_row_text:
            themes = [
                ("Semana 1", "Introducción a la Ciencia de Datos. Definición, ciclo de vida de los datos y roles.", "Actividad 'El Futuro del Científico de Datos'. Exposición grupal basada en papers (Donoho, WEF, Stanford AI). Además, taller paso a paso:\n1. Creación y configuración de cuenta en GitHub.\n2. Introducción al manejo de LLMNOTEBOOK.\n3. Uso de LLMNOTEBOOK como asistente de IA para sintetizar la exposición."),
                ("Semana 2", "Herramientas del Ecosistema. Python (Pandas, NumPy) y Jupyter Notebooks.", "Taller guiado 'Hands-on'. Programación en parejas (Pair Programming) para instalación y primeros scripts. Push inicial a GitHub. Recursos: Entorno Jupyter."),
                ("Semana 3", "Adquisición de Datos. Formatos CSV, JSON y XML.", "Búsqueda del tesoro de datos (Gamificación). Exploración de repositorios abiertos (Kaggle). Recursos: Python, Repositorios Open Data."),
                ("Semana 4", "Limpieza y Preprocesamiento. Manejo de nulos y valores atípicos.", "Aprendizaje Basado en Problemas (ABP). 'Dirty Data Challenge': limpieza de un dataset intencionalmente corrupto en equipos usando LLMNOTEBOOK como guía."),
                ("Semana 5", "Estadística Descriptiva. Medidas de tendencia central y variabilidad.", "Taller de 'Storytelling' inicial. Cálculo de medidas y justificación verbal mediante método socrático. Recursos: Dataset de negocios reales."),
                ("Semana 6", "Visualización de Datos. Histogramas, boxplots y gráficos de dispersión.", "Galería de Visualización (Data Art Walk). Co-evaluación formativa de gráficos mediante pitches de 3 minutos. Recursos: Matplotlib, Seaborn."),
                ("Semana 7", "Análisis Exploratorio de Datos (EDA). Formulación de hipótesis y correlaciones.", "Mini-Hackathon EDA. Formulación de hipótesis en tiempo real utilizando design thinking aplicado a datos. Subida de notebook a GitHub."),
                ("Semana 8", "Examen Parcial I. Evaluación de conceptos de preprocesamiento y análisis descriptivo.", "Evaluación Auténtica: Resolución de un caso práctico con límite de tiempo. Taxonomía de Bloom: Aplicar y Evaluar."),
                ("Semana 9", "Inferencia Estadística. Población, muestra y pruebas de hipótesis.", "Debate Estructurado. Discusión sobre significancia y p-value apoyado en análisis de artículos y resultados contradictorios."),
                ("Semana 10", "Aprendizaje Supervisado (Regresión). Regresión lineal simple y múltiple.", "Simulación de negocio: Predicción de demanda. Andamiaje progresivo de modelos y análisis de residuos en grupos. Recursos: Scikit-Learn."),
                ("Semana 11", "Aprendizaje Supervisado (Clasificación). Regresión logística y árboles de decisión.", "Reto 'Salva al Cliente' (Churn Prediction). Gamificación para identificar fugas de clientes. Construcción creativa de matriz de confusión."),
                ("Semana 12", "Evaluación de Modelos. Precisión, recall y validación cruzada.", "Taller 'Beat the Baseline'. Competencia formativa por optimizar un modelo base con validación cruzada. Discusión Sesgo vs Varianza."),
                ("Semana 13", "Aprendizaje No Supervisado. Agrupamiento (clustering) como K-means.", "Segmentación de Clientes. Taller analítico aplicando siluetas de agrupamiento. Presentación de perfiles descubiertos en repositorios de GitHub."),
                ("Semana 14", "Aplicaciones Especializadas. NLP, redes o series de tiempo.", "Cápsulas de conocimiento (Flipped teaching). Estudiantes investigan y presentan una demo técnica corta al grupo. Recursos: Papers técnicos."),
                ("Semana 15", "Ética y Ciencia de Datos Responsable. Privacidad y sesgos.", "Roleplay de dilemas éticos. Discusión sobre impactos sociales de algoritmos sesgados utilizando la 'Moral Machine'. Evaluación crítica."),
                ("Semana 16", "Examen Parcial II / Proyecto Final. Modelado y presentación de resultados.", "Demo Day estilo 'Shark Tank'. Evaluación por pares y feedback formativo final empleando rúbrica pública. Despliegue en GitHub.")
            ]
            
            while len(table.rows) > 1:
                table._element.remove(table.rows[1]._element)
                
            for week, content, activity in themes:
                row_cells = table.add_row().cells
                row_cells[0].text = week
                row_cells[1].text = content
                row_cells[2].text = activity

        # Evaluation Table
        if 'Actividad de evaluación' in first_row_text:
            evaluations = [
                ("Semana 8", "Evaluación Auténtica I (Preprocesamiento, EDA, Estadística descriptiva). Reto práctico situacional.", "RAE 1.1, 1.2, 1.3", "Individual", "20%"),
                ("Semanas 1-7, 9-13", "Entregas tipo Reto (Exposiciones, Laboratorios, Hackathons). Fomento de la co-evaluación formativa y uso de GitHub/LLMNOTEBOOK.", "RAE 2.1, 2.2, 2.3", "Grupal/Individual", "40%"),
                ("Semana 16", "Demo Day / Sustentación Proyecto Integrador. Evaluación con rúbrica constructiva (Andamiaje).", "RAE 3.1, 3.2, 3.3", "Grupal", "40%")
            ]
            while len(table.rows) > 1:
                table._element.remove(table.rows[1]._element)
            
            for date, activity, rae, act_type, perc in evaluations:
                row_cells = table.add_row().cells
                row_cells[0].text = date
                row_cells[1].text = activity
                row_cells[2].text = rae
                row_cells[3].text = act_type
                row_cells[4].text = perc

    doc.save(out_path)
    print(f"Document successfully created at {out_path}")

if __name__ == '__main__':
    main()
